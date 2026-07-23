from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCS_DIR = Path("/Users/phantatthanh/Documents/FPT-Sem7/KMS301/KMS301-FE/documents")

DOCX_FILES = [
    p for p in DOCS_DIR.glob("*.docx")
    if not p.name.startswith(".~") and "backup" not in p.name.lower()
]

SECTION_TITLE = "Bổ sung v1.2 - Duration metrics cho hiệu quả SOP"

BODY_INTRO = (
    "Bổ sung hai trường thời gian để đo hiệu quả vận hành của tri thức: "
    "SOP có thời gian ước tính hoàn thành làm baseline, còn bản ghi hiện trường lưu thời gian hoàn thành/thời gian thực hiện SOP thực tế. "
    "Dữ liệu này không dùng để kết luận tự động hay đánh giá cá nhân; nó là tín hiệu vòng đời tri thức để Knowledge Manager xem cùng feedback, severity, asset context, evidence và safety context."
)

ROWS = [
    [
        "SOP",
        "Thời gian ước tính hoàn thành (estimatedDurationMinutes)",
        "Contributor/KM nhập khi tạo hoặc cập nhật SOP.",
        "Tạo baseline để so sánh với thời gian thực tế khi kỹ thuật viên áp dụng SOP. Nếu actual thường xuyên chậm hơn estimate, SOP có thể thiếu bước, khó hiểu, không phù hợp context hoặc cần tách nhánh decision. Nếu actual nhanh hơn nhiều và có đề xuất hợp lý, Contributor/KM có căn cứ xem xét tối ưu SOP.",
        "ISO 9001 Clause 9 - monitoring/measurement/evaluation; ISO 30401 - review and improve KM system; KCS metrics/content health."
    ],
    [
        "Tri thức hiện trường",
        "Thời gian hoàn thành xử lý (completionTimeMinutes)",
        "Field Technician nhập trong FL-02 khi gửi case hiện trường.",
        "Ghi nhận thời gian từ lúc bắt đầu chẩn đoán/xử lý đến khi xác minh kết quả. Field này giúp nối case thực tế với baseline của SOP, hỗ trợ review lifecycle ở FL-05 và quyết định approve/reject update SOP.",
        "KCS capture-in-the-workflow and performance assessment; ISO 9001 evidence-based analysis; ISO 31000 risk/impact context."
    ],
    [
        "Áp dụng SOP",
        "Thời gian thực hiện SOP thực tế (actualDurationMinutes)",
        "Field Technician nhập khi bấm Đánh dấu đã áp dụng hoặc trong phần SOP usage của submission.",
        "Cho phép so sánh actualDurationMinutes với estimatedDurationMinutes. Đây là quantitative signal cho tri thức: liên tục chậm hơn có thể trigger review/update; nhanh hơn có thể là cơ hội cải tiến SOP nếu có evidence và proposal.",
        "KCS Evolve Loop; ISO 9001 performance evaluation; ISO 30401 continual improvement."
    ],
]

SOURCES = [
    "ISO 30401:2018 - Knowledge management systems: https://www.iso.org/standard/68683.html",
    "ISO 9001:2015 Clause 9 - Performance evaluation: https://www.iso.org/obp/ui/",
    "KCS Metrics Matrix / Performance Assessment: https://library.serviceinnovation.org/KCS/KCS_v6/KCS_v6_Practices_Guide/070",
    "KCS Content Health Indicators: https://library.serviceinnovation.org/KCS/KCS_v6/KCS_v6_Practices_Guide/030/040/010/065",
    "EPA QA/G-6 SOP Guidance: https://www.epa.gov/sites/default/files/2015-06/documents/g6-final.pdf",
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def already_has_section(document):
    return any(SECTION_TITLE in paragraph.text for paragraph in document.paragraphs)


for path in DOCX_FILES:
    doc = Document(path)
    if already_has_section(doc):
        continue

    doc.add_page_break()
    title = doc.add_heading(SECTION_TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph(BODY_INTRO)
    paragraph.style = doc.styles["Normal"]

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Phạm vi", "Field bổ sung", "Ai nhập", "Mục đích kiểm soát", "Cơ sở dẫn chứng"]
    for idx, header in enumerate(headers):
      set_cell_text(table.rows[0].cells[idx], header, bold=True)
      shade_cell(table.rows[0].cells[idx], "1F4E79")

    for row in ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if idx == 0:
                shade_cell(cells[idx], "EAF1FB")

    doc.add_paragraph("Cách dùng khi bảo vệ:")
    for bullet in [
        "Estimate time là baseline của SOP, không phải KPI phạt kỹ thuật viên.",
        "Actual/completion time là evidence thực tế để phát hiện SOP chậm, khó áp dụng hoặc có cơ hội tối ưu.",
        "Knowledge Manager dùng duration cùng evidence, feedback, severity và safety context để quyết định reconfirm, revise, suspend hoặc supersede trong FL-05.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)

    doc.add_paragraph("Nguồn tham chiếu:")
    for source in SOURCES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(source)

    doc.save(path)
    print(path)
