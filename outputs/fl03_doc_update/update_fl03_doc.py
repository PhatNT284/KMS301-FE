from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path("/Users/phantatthanh/Documents/FPT-Sem7/KMS301/KMS301-FE")
DOC_PATH = ROOT / "documents" / "Đặc_tả_Chi_tiết_FL-03_Tạo_mới_hoặc_Cập_nhật_SOP.docx"
BACKUP_PATH = ROOT / "documents" / "Đặc_tả_Chi_tiết_FL-03_Tạo_mới_hoặc_Cập_nhật_SOP.backup-before-new-sop-cta.docx"


def set_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)


def append_row_like(table, values):
    row = table.add_row()
    if len(table.rows) >= 2:
      # Reuse the previous body-row cell formatting where possible.
      template = table.rows[-2]
      for idx, cell in enumerate(row.cells):
          cell._tc.get_or_add_tcPr().append(deepcopy(template.cells[idx]._tc.get_or_add_tcPr()))
    for idx, value in enumerate(values):
        set_cell(row.cells[idx], value)
    return row


def update_matching_row(table, first_col, updates):
    for row in table.rows:
        if row.cells and row.cells[0].text.strip() == first_col:
            for idx, value in updates.items():
                set_cell(row.cells[idx], value)
            return True
    return False


def add_or_update_history(doc):
    table = doc.tables[4]
    for row in table.rows:
        if row.cells[0].text.strip() == "1.1":
            set_cell(row.cells[1], "23/07/2026")
            set_cell(row.cells[2], "Bổ sung rõ CTA Tạo SOP mới / Đề xuất SOP mới trong FL-03 để tránh lủng luồng giữa đặc tả và UI.")
            return
    append_row_like(table, [
        "1.1",
        "23/07/2026",
        "Bổ sung rõ CTA Tạo SOP mới / Đề xuất SOP mới trong FL-03 để tránh lủng luồng giữa đặc tả và UI."
    ])


def update_doc():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)

    # Document control tables and footer version.
    for table in doc.tables:
        update_matching_row(table, "Phiên bản", {1: "1.1"})
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if "Phiên bản 1.0" in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace("Phiên bản 1.0", "Phiên bản 1.1")
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Phiên bản 1.0" in cell.text:
                        set_cell(cell, cell.text.replace("Phiên bản 1.0", "Phiên bản 1.1"))

    add_or_update_history(doc)

    # Route map: current prototype route is SOP workspace, not a separate contributor-dashboard page.
    table25 = doc.tables[24]
    for row in table25.rows:
        if row.cells[0].text.strip() == "SCR-FL03-01":
            set_cell(row.cells[1], "/?screen=sops&tab=tasks hoặc /?screen=sops&tab=drafts")
            set_cell(row.cells[2], "SOP Workspace — Tasks / Drafts / New SOP CTA")
            set_cell(row.cells[5], "Task Detail; My SOP Drafts; Tạo SOP mới / Đề xuất SOP mới.")
            break

    # Screen summary for SCR-FL03-01.
    table27 = doc.tables[26]
    update_matching_row(table27, "Mục tiêu", {
        1: "Giúp Contributor nhìn thấy task SOP được giao, draft đang soạn và có một CTA rõ ràng để khởi tạo luồng SOP mới khi có business reason/source hợp lệ."
    })
    update_matching_row(table27, "Primary CTA", {
        1: "Mở task ưu tiên; Tạo SOP mới / Đề xuất SOP mới; mở My SOP Drafts."
    })

    # Component rule: make the new-SOP CTA mandatory and specify behavior.
    table28 = doc.tables[27]
    for row in table28.rows:
        if row.cells[0].text.strip() == "DASH-SOP-NEW":
            set_cell(row.cells[1], "Tạo SOP mới / Đề xuất SOP mới")
            set_cell(row.cells[2], "Secondary hoặc primary button")
            set_cell(row.cells[3], "currentUser, taxonomy, optional source")
            set_cell(row.cells[4], "CTA bắt buộc trong SOP Workspace khi role là Contributor hoặc Knowledge Manager. Bấm nút tạo NEW_SOP proposal/task hoặc draft DRAFT có business reason/source; không publish trực tiếp; route sang Task Detail hoặc Editor Step 1.")
            break

    # Add explicit action to task/detail spec if missing.
    table29 = doc.tables[28]
    if not any(row.cells[0].text.strip() == "New SOP CTA" for row in table29.rows):
        append_row_like(table29, [
            "New SOP CTA",
            "Nút Tạo SOP mới / Đề xuất SOP mới, mode=NEW_SOP, assignee=currentUser, source/businessReason.",
            "Khởi tạo proposal/task NEW_SOP; nếu tạo draft trực tiếp thì vẫn status=DRAFT và bắt buộc submit/review/publish theo SoD."
        ])

    # Business rules: make no-direct-publish and source requirement explicit.
    table64 = doc.tables[63]
    if not any(row.cells[0].text.strip() == "BR-SOP-021" for row in table64.rows):
        append_row_like(table64, [
            "BR-SOP-021",
            "CTA Tạo SOP mới phải hiển thị rõ với Contributor/Knowledge Manager trong FL-03. Hành động này chỉ tạo NEW_SOP proposal/task/draft có nguồn hoặc business reason hợp lệ; không được bỏ qua bước submit, review, SoD và publish bởi Knowledge Manager."
        ])

    # Acceptance criteria: cover the actual UI symptom found in the deployed screen.
    table72 = doc.tables[71]
    if not any(row.cells[0].text.strip().startswith("AC-FL03-24") for row in table72.rows):
        append_row_like(table72, [
            "AC-FL03-24 — Visible New SOP CTA",
            "Given user role Contributor hoặc Knowledge Manager đang ở SOP Workspace; When mở tab Nhiệm vụ SOP hoặc Draft SOP của tôi; Then phải thấy CTA Tạo SOP mới / Đề xuất SOP mới rõ ràng. When bấm CTA; Then hệ thống tạo NEW_SOP proposal/task hoặc draft DRAFT, route sang Task Detail/Editor Step 1, và vẫn yêu cầu submit + Knowledge Manager approve trước khi Published."
        ])

    # Implementation checklist: call out current route placement.
    table74 = doc.tables[73]
    if not any(row.cells[0].text.strip() == "New SOP CTA" for row in table74.rows):
        append_row_like(table74, [
            "New SOP CTA",
            "Triển khai nút Tạo SOP mới / Đề xuất SOP mới trong SopWorkspace header hoặc tab tasks/drafts; handler phải tạo NEW_SOP task/draft có traceability và tái sử dụng validation của SopEditor."
        ])

    # Add a visible inline note after the SCR-FL03-01 heading.
    note_text = "Cập nhật v1.1: Nếu UI dùng route SOP Workspace thay cho Contributor Dashboard riêng, vẫn bắt buộc hiển thị CTA Tạo SOP mới / Đề xuất SOP mới cho Contributor/Knowledge Manager. Không được chỉ dựa vào seed task NEW_SOP có sẵn vì người dùng sẽ không biết cách khởi tạo luồng mới."
    already_has_note = any(note_text in para.text for para in doc.paragraphs)
    for para in doc.paragraphs:
        if not already_has_note and para.text.strip().startswith("11.1 SCR-FL03-01"):
            p = para.insert_paragraph_before(
                note_text
            )
            p.style = para.style
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x75)
                run.bold = True
            break

    doc.save(DOC_PATH)


if __name__ == "__main__":
    update_doc()
    print(DOC_PATH)
    print(BACKUP_PATH)
